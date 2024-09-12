import os
from dotenv import load_dotenv
from sentence_transformers import SentenceTransformer
import gradio as gr
import PyPDF2
import io
from fpdf import FPDF
from typing import Union
from docx import Document
import pandas as pd
import json
import traceback
import chromadb
from chromadb.config import Settings
import numpy as np

# Initialize environment variables
load_dotenv()

# Initialize SentenceTransformer model
model = SentenceTransformer('paraphrase-MiniLM-L6-v2')

# Initialize ChromaDB client with the new configuration
chroma_client = chromadb.Client()

# Ensure the collection does not already exist
collection_name = "sentence_embeddings"
try:
    if collection_name not in [coll.name for coll in chroma_client.list_collections()]:
        collection = chroma_client.create_collection(name=collection_name)
    else:
        collection = chroma_client.get_collection(name=collection_name)
except Exception as e:
    print(f"Error initializing ChromaDB collection: {str(e)}")
    collection = None

# Function to create embeddings using SentenceTransformer
def create_embeddings(text: str):
    try:
        embeddings = model.encode([text])
        return embeddings.tolist()  # Convert ndarray to list of lists
    except Exception as e:
        print(f"Error creating embedding: {str(e)}")
        return None

# Function to add document embeddings to ChromaDB
def add_to_chromadb(text: str, doc_id: str):
    try:
        embedding = create_embeddings(text)
        if embedding is not None and collection:
            collection.add(
                documents=[text],
                embeddings=embedding,  # Should be a list of lists
                ids=[doc_id]
            )
            return "Embedding added successfully."
        else:
            return "Failed to create embedding or collection is not available."
    except Exception as e:
        return f"Error adding to ChromaDB: {str(e)}"

# Function to search embeddings with different similarity methods
def search_in_chromadb(query: str, similarity_method: str):
    try:
        embedding = create_embeddings(query)
        if embedding is not None and collection:
            # Example: Use a valid query method supported by ChromaDB
            search_result = collection.query(
                query_embeddings=embedding,  # Should be a list of lists
                n_results=5
            )
            return search_result
        else:
            return "Failed to create query embedding or collection is not available."
    except Exception as e:
        return f"Error searching in ChromaDB: {str(e)}"

# Gradio UI for embedding addition and search
def handle_document_and_search(file: Union[str, io.BytesIO], doc_type: str, query: str, similarity_method: str):
    result = ""
    if file:
        try:
            text = ""
            if isinstance(file, bytes):
                file_content = file
                if doc_type == "pdf":
                    reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                    for page in reader.pages:
                        text += page.extract_text()
                elif doc_type == "docx":
                    doc = Document(io.BytesIO(file_content))
                    text = "\n".join([para.text for para in doc.paragraphs])
                else:
                    text = file_content.decode("utf-8")
            else:
                if doc_type == "pdf":
                    reader = PyPDF2.PdfReader(file)
                    for page in reader.pages:
                        text += page.extract_text()
                elif doc_type == "docx":
                    doc = Document(file)
                    text = "\n".join([para.text for para in doc.paragraphs])
                else:
                    text = file.read().decode("utf-8")

            result = add_to_chromadb(text, doc_type + "_doc")
        except Exception as e:
            result = str(e)

    # Perform search if query is provided
    if query:
        try:
            search_result = search_in_chromadb(query, similarity_method)
            if isinstance(search_result, str):
                result += f"\n{search_result}"
            else:
                output = ""
                for doc_id, score in zip(search_result["ids"][0], search_result["distances"][0]):
                    output += f"Document ID: {doc_id}, Similarity Score: {score}\n"
                result += f"\nSearch Results:\n{output}"
        except Exception as e:
            result += f"\n{str(e)}"

    return result

# Dropdown options for similarity methods
similarity_methods = ["cosine", "dot_product", "euclidean", "manhattan"]

# Gradio Interface
with gr.Blocks() as app:
    with gr.Row():
        file_input = gr.File(label="Upload Document", type="binary")
        doc_type = gr.Radio(choices=["pdf", "docx", "txt"], label="Document Type")
        query_input = gr.Textbox(label="Search Query")
        similarity_method_input = gr.Dropdown(choices=similarity_methods, label="Similarity Method")
        action_button = gr.Button("Process")
        result_text = gr.Textbox(label="Result", lines=10)

        action_button.click(
            fn=handle_document_and_search,
            inputs=[file_input, doc_type, query_input, similarity_method_input],
            outputs=[result_text]
        )

app.launch()
