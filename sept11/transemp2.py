import os
from dotenv import load_dotenv
from sentence_transformers import SentenceTransformer
import gradio as gr
import PyPDF2
import io
from docx import Document
import numpy as np
import traceback
from typing import Union
import json
from scipy.spatial.distance import cosine
import chromadb

# Initialize environment variables
load_dotenv()

# Initialize SentenceTransformer model
model = SentenceTransformer('paraphrase-MiniLM-L6-v2')

# Initialize ChromaDB client
chroma_client = chromadb.Client()

# Ensure the collection does not already exist
collection_name = "sentence_embeddings"

def get_collection_names():
    # Retrieve list of collection names
    collections = chroma_client.list_collections()
    return [coll.name for coll in collections]

if collection_name not in get_collection_names():
    collection = chroma_client.create_collection(name=collection_name)
else:
    collection = chroma_client.get_collection(name=collection_name)

# Function to create embeddings using SentenceTransformer
def create_embeddings(text: str):
    try:
        # Encode the text using SentenceTransformer
        embeddings = model.encode([text])
        return embeddings[0]  # Return the first (and only) embedding
    except Exception as e:
        print(f"Error creating embedding: {str(e)}")
        return None

# Function to add document embeddings to ChromaDB
def add_to_chromadb(text: str, doc_id: str):
    embedding = create_embeddings(text)
    if embedding is not None:
        collection.add(
            documents=[text],
            embeddings=[embedding.tolist()],  # Convert to list for compatibility
            ids=[doc_id]
        )
        return "Embedding added successfully."
    else:
        return "Failed to create embedding."

# Function to get all embeddings from ChromaDB
def get_all_embeddings():
    docs = collection.get()
    return docs['embeddings']

# Function to search embeddings with similarity calculation
def search(query: str):
    try:
        query_embedding = create_embeddings(query)
        if query_embedding is None:
            return "Failed to create query embedding."

        # Retrieve all stored embeddings
        all_embeddings = get_all_embeddings()
        
        # Compute similarities
        similarities = []
        for doc_id, embedding in all_embeddings.items():
            # Compute cosine similarity
            similarity = 1 - cosine(query_embedding, embedding)
            similarities.append((doc_id, similarity))

        # Sort by similarity score in descending order
        similarities.sort(key=lambda x: x[1], reverse=True)
        return similarities[:5]  # Return top 5 results
    except Exception as e:
        traceback.print_exc()
        return str(e)

# Gradio UI for embedding addition and search
def add_document_to_collection(file: Union[str, io.BytesIO], doc_type: str):
    try:
        text = ""
        if isinstance(file, bytes):
            # Handle bytes directly
            file_content = file
            if doc_type == "pdf":
                reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                for page in reader.pages:
                    text += page.extract_text()
            elif doc_type == "docx":
                doc = Document(io.BytesIO(file_content))
                text = "\n".join([para.text for para in doc.paragraphs])
            else:
                text = file_content.decode("utf-8")
        else:
            # Handle file-like object
            if doc_type == "pdf":
                reader = PyPDF2.PdfReader(file)
                for page in reader.pages:
                    text += page.extract_text()
            elif doc_type == "docx":
                doc = Document(file)
                text = "\n".join([para.text for para in doc.paragraphs])
            else:
                text = file.read().decode("utf-8")
        
        result = add_to_chromadb(text, doc_type + "_doc")
        return result
    except Exception as e:
        traceback.print_exc()
        return str(e)


# Gradio UI for searching
def search_documents(query: str):
    results = search(query)
    if isinstance(results, str):
        return results
    output = ""
    for doc_id, score in results:
        output += f"Document ID: {doc_id}, Similarity Score: {score}\n"
    return output

# Gradio Interface
with gr.Blocks() as app:
    with gr.Tab("Add Document"):
        file_input = gr.File(label="Upload Document", type="binary")
        doc_type = gr.Radio(choices=["pdf", "docx", "txt"], label="Document Type")
        add_button = gr.Button("Add to ChromaDB")
        result_text = gr.Textbox(label="Result")

        add_button.click(
            fn=add_document_to_collection, 
            inputs=[file_input, doc_type], 
            outputs=[result_text]
        )
    
    with gr.Tab("Search"):
        query_input = gr.Textbox(label="Search Query")
        search_button = gr.Button("Search")
        search_results = gr.Textbox(label="Search Results")

        search_button.click(
            fn=search_documents, 
            inputs=[query_input], 
            outputs=[search_results]
        )

app.launch()
