import os
from dotenv import load_dotenv
from sentence_transformers import SentenceTransformer
import gradio as gr
import PyPDF2
import io
from fpdf import FPDF
from typing import Union
from docx import Document
import numpy as np
from scipy.spatial.distance import cosine, euclidean, cityblock
import chromadb

# Initialize environment variables
load_dotenv()

# Initialize SentenceTransformer model
model = SentenceTransformer('paraphrase-MiniLM-L6-v2')

# Initialize ChromaDB client
chroma_client = chromadb.Client()

# Ensure the collection does not already exist
collection_name = "sentence_embeddings"
try:
    if collection_name not in [coll.name for coll in chroma_client.list_collections()]:
        collection = chroma_client.create_collection(name=collection_name)
    else:
        collection = chroma_client.get_collection(name=collection_name)
except Exception as e:
    print(f"Error initializing ChromaDB collection: {str(e)}")
    collection = None

# Function to create embeddings using SentenceTransformer
def create_embeddings(text: str):
    try:
        embeddings = model.encode([text])
        return embeddings[0].tolist()  # Convert ndarray to list
    except Exception as e:
        print(f"Error creating embedding: {str(e)}")
        return None

# Function to add document embeddings to ChromaDB
def add_to_chromadb(text: str, doc_id: str):
    try:
        embedding = create_embeddings(text)
        if embedding is not None and collection:
            collection.add(
                documents=[text],
                embeddings=[embedding],  # Should be a list of lists
                ids=[doc_id]
            )
            return "Embedding added successfully."
        else:
            return "Failed to create embedding or collection is not available."
    except Exception as e:
        return f"Error adding to ChromaDB: {str(e)}"

# Function to calculate similarity metrics
def calculate_similarity(query_embedding, doc_embedding, similarity_method):
    if similarity_method == "cosine":
        return 1 - cosine(query_embedding, doc_embedding)
    elif similarity_method == "dot_product":
        return np.dot(query_embedding, doc_embedding)
    elif similarity_method == "euclidean":
        return -euclidean(query_embedding, doc_embedding)
    elif similarity_method == "manhattan":
        return -cityblock(query_embedding, doc_embedding)
    else:
        raise ValueError("Unsupported similarity method")

# Function to search embeddings with different similarity methods
def search_in_chromadb(query: str, similarity_method: str):
    try:
        query_embedding = create_embeddings(query)
        if query_embedding is not None and collection:
            results = collection.query(
                query_embeddings=[query_embedding],
                n_results=10  # Adjust based on how many results you want
            )
            search_results = results["results"]
            output = []
            for result in search_results:
                doc_id = result["id"]
                score = result["score"]
                doc_embedding = result["embedding"]  # Retrieve stored embedding
                similarity_score = calculate_similarity(query_embedding, doc_embedding, similarity_method)
                output.append((doc_id, similarity_score, doc_embedding))

            # Sort results by similarity score
            output.sort(key=lambda x: x[1], reverse=True)
            return output
        else:
            return "Failed to create query embedding or collection is not available."
    except Exception as e:
        return f"Error searching in ChromaDB: {str(e)}"

# Function to handle document and search
def handle_document_and_search(file: Union[str, io.BytesIO], doc_type: str, query: str, similarity_method: str):
    result = ""
    if file:
        try:
            text = ""
            if isinstance(file, bytes):
                file_content = file
                if doc_type == "pdf":
                    reader = PyPDF2.PdfReader(io.BytesIO(file_content))
                    for page in reader.pages:
                        text += page.extract_text()
                elif doc_type == "docx":
                    doc = Document(io.BytesIO(file_content))
                    text = "\n".join([para.text for para in doc.paragraphs])
                else:
                    text = file_content.decode("utf-8")
            else:
                if doc_type == "pdf":
                    reader = PyPDF2.PdfReader(file)
                    for page in reader.pages:
                        text += page.extract_text()
                elif doc_type == "docx":
                    doc = Document(file)
                    text = "\n".join([para.text for para in doc.paragraphs])
                else:
                    text = file.read().decode("utf-8")

            result = add_to_chromadb(text, doc_type + "_doc")
        except Exception as e:
            result = str(e)

    # Perform search if query is provided
    if query:
        try:
            search_result = search_in_chromadb(query, similarity_method)
            if isinstance(search_result, str):
                result += f"\n{search_result}"
            else:
                output = ""
                for doc_id, score, embedding in search_result:
                    output += f"Document ID: {doc_id}, Similarity Score: {score}, Embedding: {embedding}\n"
                result += f"\nSearch Results:\n{output}"
        except Exception as e:
            result += f"\n{str(e)}"

    return result

# Dropdown options for similarity methods
similarity_methods = ["cosine", "dot_product", "euclidean", "manhattan"]

# Gradio Interface
with gr.Blocks() as app:
    with gr.Row():
        file_input = gr.File(label="Upload Document", type="binary")
        doc_type = gr.Radio(choices=["pdf", "docx", "txt"], label="Document Type")
        query_input = gr.Textbox(label="Search Query")
        similarity_method_input = gr.Dropdown(choices=similarity_methods, label="Similarity Method")
        action_button = gr.Button("Process")
        result_text = gr.Textbox(label="Result", lines=15)

        action_button.click(
            fn=handle_document_and_search,
            inputs=[file_input, doc_type, query_input, similarity_method_input],
            outputs=[result_text]
        )

app.launch()
